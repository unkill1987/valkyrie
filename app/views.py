
import requests
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect
import hashlib
from docx import Document
from docx.shared import Cm
import os
import time

from app.models import Contract_LCR, Member, Contract_CI, Contract_SR, Contract_BL, Contract_DO, Contract_LC,Process, Process_complete
from valweb import settings
from django.utils import timezone

def about(request):
    return render(request,'app/about.html',{})

def search(request):
    try:
        cid = str(request.POST['cid'])
        url = ("http://222.239.231.247:8001/keyHistory/%s" % cid)
        res = requests.post(url)
        history = res.json()
        if len(history) == 0:
            message = "Invalid Contract ID"
            print(message)
            return render(request, 'app/index.html', {'message': message})
        else:
            history.reverse()
            return render(request, 'app/search.html', {'cid': cid, 'history': history})
    except Exception as e:
        print(e)
        return redirect('index')


def share1(request):

    try:
        check_id = request.POST['check_id']
        share_user = request.POST['share_user']
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LCR.objects.filter(owner=member)
        getid = contract.filter(id=check_id)
        hash = getid.values('sha256')[0]['sha256']
        contract_id = getid.values('contract_id')[0]['contract_id']
        contract = contract_id
        user_id = getid.values('owner')[0]['owner']


        url = ('http://222.239.231.247:8001/add_LCR/' + contract + '-' + user_id + '-' + hash)
        response = requests.post(url)
        res = response.text
        result_dict = {}

        if (res == "The contract already exists"):
           result_dict['result'] = 'Fail'

        else:
            share = Contract_LCR.objects.get(id=check_id)
            share.share3 = share_user
            share.save()

            process = Process.objects.get(id=contract_id)
            process.user1 = user_id
            process.LCR_hash = hash
            process.user3 = share_user
            process.save()

            process_complete = Process_complete.objects.get(id=contract_id)
            process_complete.LCR_hash = hash
            process_complete.save()
            result_dict['result'] = 'Success'
        return JsonResponse(result_dict)
    except Exception as e:
        print(e)
        return redirect('ing')


def share2(request):


    try:
        check_id = request.POST['check_id']
        share_user = request.POST['share_user']
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_CI.objects.filter(owner=member)
        getid = contract.filter(id=check_id)
        hash = getid.values('sha256')[0]['sha256']
        contract_id = getid.values('id')[0]['id']
        contract = str(contract_id)
        user_id = getid.values('owner')[0]['owner']


        url = 'http://222.239.231.247:8001/add_CI/'+ contract +'-'+user_id+'-'+hash
        response = requests.post(url)
        res = response.text
        result_dict = {}

        if ( res == "The contract already exists" ):
            result_dict['result'] = 'Fail'
        else:
            share = Contract_CI.objects.get(id=check_id)
            share.share1 = share_user
            share.save()

            process = Process.objects.get(CI_hash=hash)
            process.user1 = share_user
            process.save()
            result_dict['result'] = 'Success'
        return JsonResponse(result_dict)
    except Exception as e:
        print(e)
        return redirect('ing2')

def share2_1(request):


    try:
        check_id = request.POST['check_id']
        share_user = request.POST['share_user']
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_SR.objects.filter(owner=member)
        getid = contract.filter(id=check_id)
        hash = getid.values('sha256')[0]['sha256']
        contract_id = getid.values('contract_id')[0]['contract_id']
        contract = str(contract_id)
        user_id = getid.values('owner')[0]['owner']


        url = 'http://222.239.231.247:8001/add_SR/' + contract + '-' + user_id + '-' + hash
        response = requests.post(url)
        res = response.text
        result_dict = {}

        if (res == "The contract already exists"):
            result_dict['result'] = 'Fail'
        else:
            process = Process.objects.get(id=contract_id)
            process.SR_hash = hash
            process.user4 = share_user
            process.save()

            process_complete = Process_complete.objects.get(id=contract_id)
            process_complete.SR_hash = hash
            process_complete.save()

            share = Contract_SR.objects.get(id=check_id)
            share.share4 = share_user
            share.save()
            result_dict['result'] = 'Success'
        return JsonResponse(result_dict)
    except Exception as e:
        print (e)
        return redirect('ing2_1')

def share3(request):


    try:
        check_id = request.POST['check_id']
        share_user = request.POST['share_user']
        share_user2 = request.POST['share_user2']
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LC.objects.filter(owner=member)
        getid = contract.filter(id=check_id)
        hash = getid.values('sha256')[0]['sha256']
        contract_id = getid.values('contract_id')[0]['contract_id']
        contract = str(contract_id)
        user_id = getid.values('owner')[0]['owner']


        url = 'http://222.239.231.247:8001/add_LC/' + contract + '-' + user_id + '-' + hash
        response = requests.post(url)
        res = response.text
        result_dict = {}

        if (res == "The contract already exists"):
            result_dict['result'] = 'Fail'
        else:
            share = Contract_LC.objects.get(id=check_id)
            share.share1 = share_user
            share.share2 = share_user2
            share.save()

            process = Process.objects.get(id=contract_id)
            process.LC_hash = hash
            process.save()

            process_complete = Process_complete.objects.get(id=contract_id)
            process_complete.LC_hash = hash
            process_complete.save()
            result_dict['result'] = 'Success'
        return JsonResponse(result_dict)
    except Exception as e:
        print(e)
        return redirect('ing3')

def share4_1(request):


    try:
        check_id = request.POST['check_id']
        share_user = request.POST['share_user']
        share_user2 = request.POST['share_user2']
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_BL.objects.filter(owner=member)
        getid = contract.filter(id=check_id)
        hash = getid.values('sha256')[0]['sha256']
        contract_id = getid.values('contract_id')[0]['contract_id']
        contract = str(contract_id)
        user_id = getid.values('owner')[0]['owner']


        url = 'http://222.239.231.247:8001/add_BL/' + contract + '-' + user_id + '-' + hash
        response = requests.post(url)
        res = response.text
        result_dict = {}

        if (res == "The contract already exists"):
            result_dict['result'] = 'Fail'
        else:
            share = Contract_BL.objects.get(id=check_id)
            share.share1 = share_user
            share.share2 = share_user2
            share.save()

            process = Process.objects.get(id=contract_id)
            process.BL_hash = hash
            process.save()

            process_complete = Process_complete.objects.get(id=contract_id)
            process_complete.BL_hash = hash
            process_complete.save()

            result_dict['result'] = 'Success'
        return JsonResponse(result_dict)
    except:
        return redirect('ing4_1')

def share4_2(request):


    try:
        check_id = request.POST['check_id']
        share_user = request.POST['share_user']
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_DO.objects.filter(owner=member)
        getid = contract.filter(id=check_id)
        hash = getid.values('sha256')[0]['sha256']
        contract_id = getid.values('contract_id')[0]['contract_id']
        contract = str(contract_id)
        user_id = getid.values('owner')[0]['owner']


        url = 'http://222.239.231.247:8001/add_DO/' + contract + '-' + user_id + '-' + hash
        response = requests.post(url)
        res = response.text
        result_dict = {}

        if (res == "The contract already exists"):
            result_dict['result'] = 'Fail'
        else:
            share = Contract_DO.objects.get(id=check_id)
            share.share1 = share_user
            share.save()

            process = Process.objects.get(id=contract_id)
            process.DO_hash = hash
            process.save()

            process_complete = Process_complete.objects.get(id=contract_id)
            process_complete.DO_hash = hash
            process_complete.save()

            result_dict['result'] = 'Success'
        return JsonResponse(result_dict)
    except:
        return redirect('ing4_2')


def remove(request):


    result_dict={}
    try:
        check_id = request.POST['check_id']
        Contract_LCR.objects.get(id=check_id).delete()
        result_dict['result'] = 'Deleted'
        return JsonResponse(result_dict)
    except:
        return redirect('ing')

def remove2(request):
    check_id = request.POST['check_id']
    result_dict = {}

    try:
        Contract_CI.objects.get(id=check_id).delete()
        result_dict['result'] = 'Deleted'
        return JsonResponse(result_dict)
    except:
        return redirect('ing2')

def remove2_1(request):

    result_dict = {}

    try:
        check_id = request.POST['check_id']
        Contract_SR.objects.get(id=check_id).delete()
        result_dict['result'] = 'Deleted'
        return JsonResponse(result_dict)
    except:
        return redirect('ing2_1')

def remove3(request):

    result_dict = {}

    try:
        check_id = request.POST['check_id']
        Contract_LC.objects.get(id=check_id).delete()
        result_dict['result'] = 'Deleted'
        return JsonResponse(result_dict)
    except:
        return redirect('ing3')



def remove4_1(request):

    result_dict = {}

    try:
        check_id = request.POST['check_id']
        Contract_BL.objects.get(id=check_id).delete()
        result_dict['result'] = 'Deleted'
        return JsonResponse(result_dict)
    except:
        return redirect('ing4_1')

def remove4_2(request):

    result_dict = {}

    try:
        check_id = request.POST['check_id']
        Contract_DO.objects.get(id=check_id).delete()
        result_dict['result'] = 'Deleted'
        return JsonResponse(result_dict)
    except:
        return redirect('ing4_2')


def process1_remove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            remove = Process.objects.get(id=id)
            remove.user1 = ' '
            remove.save()
        except Exception as e:
            print(e)
            pass

    return redirect('process1')

def process2_remove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            remove = Process.objects.get(id=id)
            remove.user2 = ' '
            remove.save()
        except Exception as e:
            print(e)
            pass

    return redirect('process2')

def process3_remove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            remove = Process.objects.get(id=id)
            remove.user3 = ' '
            remove.save()
        except Exception as e:
            print(e)
            pass

    return redirect('process3')

def process4_remove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            remove = Process.objects.get(id=id)
            remove.user4 = ' '
            remove.save()
        except Exception as e:
            print(e)
            pass

    return redirect('process4')


def ciremove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_CI.objects.get(id=id)
            share.share1 = ' '
            share.save()
        except:
            pass
    return redirect('cireceived')


def blremove1(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_BL.objects.get(id=id)
            share.share1 = ' '
            share.save()
        except:
            pass
    return redirect('blreceived1')

def lcremove1(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_LC.objects.get(id=id)
            share.share1 = ' '
            share.save()
        except:
            pass
    return redirect('lcreceived1')

def blremove2(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_BL.objects.get(id=id)
            share.share2 = ' '
            share.save()
        except:
            pass
    return redirect('blreceived2')


def lcremove2(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_LC.objects.get(id=id)
            share.share2 = ' '
            share.save()
        except:
            pass
    return redirect('lcreceived2')

def doremove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_DO.objects.get(id=id)
            share.share2 = ' '
            share.save()
        except:
            pass
    return redirect('doreceived')

def lcrremove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_LCR.objects.get(id=id)
            share.share3 = ' '
            share.save()
        except:
            pass
    return redirect('lcrreceived')


def srremove(request):
    # deletes all objects from Car database table
    # Contract.objects.get('id').delete()
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            share = Contract_SR.objects.get(id=id)
            share.share4 = ' '
            share.save()
        except:
            pass
    return redirect('cireceived')

def process1_complete(request):
    user_id= request.session['user_id']
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            process = Process.objects.get(id=id)
            process.user1 = ' '
            process.save()
            process_complete = Process_complete.objects.get(id=id)
            process_complete.user1 = user_id
            process_complete.save()
        except:
            pass
    return redirect('process1')

def process2_complete(request):
    user_id= request.session['user_id']
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            process = Process.objects.get(id=id)
            process.user2 = ' '
            process.save()
            process_complete = Process_complete.objects.get(id=id)
            process_complete.user2 = user_id
            process_complete.save()
        except Exception as e:
            print (e)
            pass
    return redirect('process2')

def process3_complete(request):
    user_id= request.session['user_id']
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            process = Process.objects.get(id=id)
            process.user3 = ' '
            process.save()
            process_complete = Process_complete.objects.get(id=id)
            process_complete.user3 = user_id
            process_complete.save()
        except:
            pass
    return redirect('process3')

def process4_complete(request):
    user_id= request.session['user_id']
    check_id = request.GET['check_id']
    check_ids = check_id.split(',')

    for id in check_ids:
        try:
            process = Process.objects.get(id=id)
            process.user4 = ' '
            process.save()
            process_complete = Process_complete.objects.get(id=id)
            process_complete.user4 = user_id
            process_complete.save()
        except Exception as e:
            print(e)
            pass
    return redirect('process4')



def submit(request):
    contractname = request.POST['contractname']
    contract_id = request.POST['contract_id']
    a = request.POST['a']
    b = request.POST['b']
    c = request.POST['c']
    d = request.POST['d']
    e = request.POST['e']
    f = request.POST['f']
    g = request.POST['g']
    h = request.POST['h']
    i = request.POST['i']
    j = request.POST['j']
    k = request.POST['k']
    l = request.POST['l']
    m = request.POST['m']
    n = request.POST['n']
    o = request.POST['o']
    time_format = time.strftime('%Y-%m-%d_%H%M%S', time.localtime(time.time()))


    records = (
        (1, 'Advising bank:', a),
        (2, 'Credit No.:', b),
        (3, 'Beneficiary:', c),
        (4, 'Applicant:', d),
        (5, 'L/C Amount and Tolerance:', e),
        (6, 'Type:', f),
        (7, 'Partial shipment:', g),
        (8, 'Transshipment:', h),
        (9, 'Trnasport mode:', i),
        (10, 'Loading(shipment from):', j),
        (11, 'Discharging(shipment to):', k),
        (12, 'Latest shipment date:', l),
        (13, 'All banking charges:', m),
        (14, 'Confirmation:', n),
        (15, 'T/T reimbursement:', o),
    )

    document = Document()
    document.add_heading(contractname + '/' +time_format, 0)

    table = document.add_table(rows = 1, cols =3)


    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.5)
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Content'

    for num, title, content in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = title
        row_cells[2].text = content

    document.save('document/LCR_'+ time_format + '.docx')

    file = open('document/LCR_' + time_format + '.docx', 'rb')
    data = file.read()

    hash = hashlib.sha256(data).hexdigest()
    file.close()


    # 데이터 저장
    contract = Contract_LCR(contractname=contractname, contract_id = contract_id, sha256=hash, filename='document/LCR_' + time_format + '.docx')

    # 로그인한 사용자 정보를 Contract에 같이 저장
    user_id = request.session['user_id']
    member = Member.objects.get(user_id=user_id)
    contract.owner = member
    contract.save()
    return redirect('ing')



def submit2(request):
    invoicename = request.POST['invoicename']
    a = request.POST['a']
    b = request.POST['b']
    c = request.POST['c']
    d = request.POST['d']
    e = request.POST['e']
    f = request.POST['f']
    g = request.POST['g']
    h = request.POST['h']
    i = request.POST['i']
    j = request.POST['j']
    k = request.POST['k']
    l = request.POST['l']
    m = request.POST['m']
    n = request.POST['n']
    o = request.POST['o']
    p = request.POST['p']
    q = request.POST['q']
    r = request.POST['r']

    time_format = time.strftime('%Y-%m-%d_%H%M%S', time.localtime(time.time()))

    records = (
        (1, 'Shipper/Seller:', a),
        (2, 'Consignee:', b),
        (3, 'Departure Date:', c),
        (4, 'Vessel/Flight:', d),
        (5, 'From:', e),
        (6, 'To:', f),
        (7, 'Invoice No.and Date:', g),
        (8, 'L/C No.and Date:', h),
        (9, 'Buyer(if other than consignee):', i),
        (10, 'Other reference:', j),
        (11, 'Terms of delivery and payment:', k),
        (12, 'Shipping Mark:', l),
        (13, 'No.and kind of packages:', m),
        (14, 'Goods description:', n),
        (15, 'Quantity:', o),
        (16, 'Unit price:', p),
        (17, 'Amount:', q),
        (18, 'Singed by:', r),
    )

    document = Document()
    document.add_heading(invoicename + '/' + time_format, 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.5)
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Content'

    for num, title, content in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = title
        row_cells[2].text = content

    document.save('document/CI_' + time_format + '.docx')

    file = open('document/CI_' + time_format + '.docx', 'rb')
    data = file.read()

    hash = hashlib.sha256(data).hexdigest()
    file.close()

    # 데이터 저장
    contract = Contract_CI(contractname=invoicename, sha256=hash, filename='document/CI_' + time_format + '.docx')

    # 로그인한 사용자 정보를 Contract에 같이 저장
    user_id = request.session['user_id']
    member = Member.objects.get(user_id=user_id)
    contract.owner = member

    contract.save()
    id = Contract_CI.objects.filter(sha256=hash).values('id')[0]['id']

    process = Process(contract_id=id, user2=user_id, CI_hash=hash)
    process.save()
    process_complete = Process_complete(contract_id=id, CI_hash=hash)
    process_complete.save()

    return redirect('ing2')


def submit2_1(request):
    srname = request.POST['srequestname']
    contract_id = request.POST['contract_id']
    a = request.POST['a']
    b = request.POST['b']
    c = request.POST['c']
    d = request.POST['d']
    e = request.POST['e']
    f = request.POST['f']
    g = request.POST['g']
    h = request.POST['h']
    i = request.POST['i']
    j = request.POST['j']
    k = request.POST['k']
    l = request.POST['l']
    m = request.POST['m']
    n = request.POST['n']
    o = request.POST['o']
    time_format = time.strftime('%Y-%m-%d_%H%M%S', time.localtime(time.time()))


    records = (
        (1, 'Shipper:', a),
        (2, 'Consignee:', b),
        (3, 'Notify Party:', c),
        (4, 'Vessel:', d),
        (5, 'Voyage No.:', e),
        (6, 'Port of Loading:', f),
        (7, 'Port of Discharge:', g),
        (8, 'Final Destination:', h),
        (9, 'Marking:', i),
        (10, 'Packages:', j),
        (11, 'Description of Goods:', k),
        (12, 'Gross Weight:', l),
        (13, 'Measurement:', m),
        (14, 'Freight Term:', n),
        (15, 'Original B/L:', o),

    )

    document = Document()
    document.add_heading(srname + '/' + time_format, 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.5)
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Content'

    for num, title, content in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = title
        row_cells[2].text = content

    document.save('document/SR_' + time_format + '.docx')

    file = open('document/SR_' + time_format + '.docx', 'rb')
    data = file.read()

    hash = hashlib.sha256(data).hexdigest()
    file.close()

    # 데이터 저장
    contract = Contract_SR(contractname=srname, contract_id=contract_id, sha256=hash, filename='document/SR_' + time_format + '.docx')

    # 로그인한 사용자 정보를 Contract에 같이 저장
    user_id = request.session['user_id']
    member = Member.objects.get(user_id=user_id)
    contract.owner = member

    contract.save()

    return redirect('ing2_1')


def submit3(request):
    letteroflc = request.POST['letteroflc']
    contract_id = request.POST['contract_id']
    a = request.POST['a']
    b = request.POST['b']
    c = request.POST['c']
    d = request.POST['d']
    e = request.POST['e']
    f = request.POST['f']
    g = request.POST['g']
    h = request.POST['h']
    i = request.POST['i']
    j = request.POST['j']
    k = request.POST['k']
    l = request.POST['l']
    m = request.POST['m']
    n = request.POST['n']
    o = request.POST['o']
    p = request.POST['p']
    q = request.POST['q']
    r = request.POST['r']
    s = request.POST['s']
    t = request.POST['t']
    u = request.POST['u']
    v = request.POST['v']
    w = request.POST['w']
    x = request.POST['x']
    y = request.POST['y']
    z = request.POST['z']
    aa = request.POST['aa']
    bb = request.POST['bb']
    cc = request.POST['cc']
    dd = request.POST['dd']
    ee = request.POST['ee']
    ff = request.POST['ff']
    gg = request.POST['gg']

    time_format = time.strftime('%Y-%m-%d_%H%M%S', time.localtime(time.time()))

    records = (
        (1, 'Transfer:', a),
        (2, 'Credit Number:', b),
        (3, 'Advising Bank:', c),
        (4, 'Expiry Date:', d),
        (5, 'Applicant:', e),
        (6, 'Beneficiary:', f),
        (7, 'Amount:', g),
        (8, 'Partial Shipment:', h),
        (9, 'Latest Shipment Date:', i),
        (10, 'Additional Conditions:', j),
        (11, 'All banking charges:', k),
        (12, 'Documents delivered by:', l),
        (13, 'Confirmation:', m),
        (14, 'Reissue:', n),
        (15, 'Import L/C Transfer:', o),
        (16, 'Draft at:', p),
        (17, 'Usance:', q),
        (18, 'SettlingBank:', r),
        (19, 'Credit:', s),
        (20, 'Transshipment mode:', t),
        (21, 'Authorization:', u),
        (22, 'Port of Loading/Airport of Departure:', v),
        (23, 'Place of Taking in Charge:', w),
        (24, 'Signed/Original/Commercial Invoice:', x),
        (25, 'Full Set of B/L', y),
        (26, 'Certificate of Origin in:', z),
        (27, 'Certificate of Analysis in:', aa),
        (28, 'Other Documents Required:', bb),
        (29, 'Description of Goods/Services:', cc),
        (30, 'Price Terms:', dd),
        (31, 'Country of Origin:', ee),
        (32, 'HS CODE:', ff),
        (33, 'CommodityDescription:', gg),

    )

    document = Document()
    document.add_heading(letteroflc + '/' + time_format, 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.5)
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Content'

    for num, title, content in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = title
        row_cells[2].text = content

    document.save('document/LC_' + time_format + '.docx')

    file = open('document/LC_' + time_format + '.docx', 'rb')
    data = file.read()

    hash = hashlib.sha256(data).hexdigest()
    file.close()
    # 데이터 저장
    contract = Contract_LC(contractname=letteroflc, contract_id=contract_id, sha256=hash, filename='document/LC_' + time_format + '.docx')

    # 로그인한 사용자 정보를 Contract에 같이 저장
    user_id = request.session['user_id']
    member = Member.objects.get(user_id=user_id)
    contract.owner = member

    contract.save()


    return redirect('ing3')


def submit4_1(request):
    contractname = request.POST['contractname']
    contract_id = request.POST['contract_id']
    a = request.POST['a']
    b = request.POST['b']
    c = request.POST['c']
    d = request.POST['d']
    e = request.POST['e']
    f = request.POST['f']
    g = request.POST['g']
    h = request.POST['h']
    i = request.POST['i']
    j = request.POST['j']

    time_format = time.strftime('%Y-%m-%d_%H%M%S', time.localtime(time.time()))

    records = (
        (1, 'Bank:', a),
        (2, 'Nodify party:', b),
        (3, 'Vessel:', c),
        (4, 'Voyage No.:', d),
        (5, 'Part of loading:', e),
        (6, 'Place of receipt:', f),
        (7, 'Place of delivery:', g),
        (8, 'Description of goods:', h),
        (9, 'Weight:', i),
        (10, 'Measurement:', j),

    )

    document = Document()
    document.add_heading(contractname + '/' + time_format, 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.5)
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Content'

    for num, title, content in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = title
        row_cells[2].text = content

    document.save('document/BL_' + time_format + '.docx')

    file = open('document/BL_' + time_format + '.docx', 'rb')
    data = file.read()

    hash = hashlib.sha256(data).hexdigest()
    file.close()

    # 데이터 저장
    contract = Contract_BL(contractname=contractname, contract_id=contract_id, sha256=hash, filename='document/BL_' + time_format + '.docx')

    # 로그인한 사용자 정보를 Contract에 같이 저장
    user_id = request.session['user_id']
    member = Member.objects.get(user_id=user_id)
    contract.owner = member

    contract.save()

    return redirect('ing4_1')


def submit4_2(request):
    contractname = request.POST['contractname']
    contract_id = request.POST['contract_id']
    a = request.POST['a']
    b = request.POST['b']
    c = request.POST['c']
    d = request.POST['d']
    e = request.POST['e']
    f = request.POST['f']
    g = request.POST['g']

    time_format = time.strftime('%Y-%m-%d_%H%M%S', time.localtime(time.time()))

    records = (
        (1, 'Agent name:', a),
        (2, 'Restricted delivery(Yes or no):', b),
        (3, 'Adult signature restriced delivery(Yes or no):', c),
        (4, 'Agent Signture:', d),
        (5, 'ID verified (yes or no):', e),
        (6, 'USPS initals:', f),
        (7, 'Date:', g),

    )

    document = Document()
    document.add_heading(contractname + '/' + time_format, 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.5)
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Content'

    for num, title, content in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = title
        row_cells[2].text = content

    document.save('document/DO_' + time_format + '.docx')

    file = open('document/DO_' + time_format + '.docx', 'rb')
    data = file.read()

    hash = hashlib.sha256(data).hexdigest()
    file.close()
    # 데이터 저장
    contract = Contract_DO(contractname=contractname, contract_id=contract_id, sha256=hash, filename='document/DO_' + time_format + '.docx')

    # 로그인한 사용자 정보를 Contract에 같이 저장
    user_id = request.session['user_id']
    member = Member.objects.get(user_id=user_id)
    contract.owner = member
    contract.save()

    return redirect('ing4_2')

def download(request):
    id = request.GET['id']
    c = Contract_LCR.objects.get(id=id)

    filepath = os.path.join(settings.BASE_DIR, c.filename)
    # filepath = os.path.join('/home/valkyrie1234', c.filename)
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as f:
        response = HttpResponse(f, content_type='text/plain')
        response['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        return response


def download2(request):
    id = request.GET['id']
    c = Contract_CI.objects.get(id=id)

    filepath = os.path.join(settings.BASE_DIR, c.filename)
    # filepath = os.path.join('/home/valkyrie1234', c.filename)
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as f:
        response = HttpResponse(f, content_type='text/plain')
        response['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        return response

def download2_1(request):
    id = request.GET['id']
    c = Contract_SR.objects.get(id=id)

    filepath = os.path.join(settings.BASE_DIR, c.filename)
    # filepath = os.path.join('/home/valkyrie1234', c.filename)
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as f:
        response = HttpResponse(f, content_type='text/plain')
        response['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        return response

def download3(request):
    id = request.GET['id']
    c = Contract_LC.objects.get(id=id)

    filepath = os.path.join(settings.BASE_DIR, c.filename)
    # filepath = os.path.join('/home/valkyrie1234', c.filename)
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as f:
        response = HttpResponse(f, content_type='text/plain')
        response['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        return response

def download4_1(request):
    id = request.GET['id']
    c = Contract_BL.objects.get(id=id)

    filepath = os.path.join(settings.BASE_DIR, c.filename)
    # filepath = os.path.join('/home/valkyrie1234', c.filename)
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as f:
        response = HttpResponse(f, content_type='text/plain')
        response['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        return response

def download4_2(request):
    id = request.GET['id']
    c = Contract_DO.objects.get(id=id)
    print(c.filename)

    filepath = os.path.join(settings.BASE_DIR, c.filename)
    # filepath = os.path.join('/home/valkyrie1234', c.filename)
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as f:
        response = HttpResponse(f, content_type='text/plain')
        response['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        return response

def process1(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        try:
            contract = Process.objects.filter(user1=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract=None
            n=0

        return render(request, 'app/process1.html',{'contract':contracts, 'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')

def process2(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process.objects.filter(user2=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process2.html', {'contract': contracts,'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')

def process3(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process.objects.filter(user3=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process3.html', {'contract': contracts, 'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')

def process4(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process.objects.filter(user4=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process4.html', {'contract': contracts,'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')

def process1_done(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process_complete.objects.filter(user1=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process1_done.html', {'contract': contracts,'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')

def process2_done(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process_complete.objects.filter(user2=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process2_done.html', {'contract': contracts,'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')


def process3_done(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process_complete.objects.filter(user3=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process3_done.html', {'contract': contracts,'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')


def process4_done(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)

        try:
            contract = Process_complete.objects.filter(user4=member).order_by('-id')
            n = len(contract)
            paginator = Paginator(contract, 3)
            page = request.GET.get('page')
            contracts = paginator.get_page(page)

        except:
            contract = None
            n = 0

        return render(request, 'app/process4_done.html', {'contract': contracts,'n':n})
    except Exception as e:
        print(e)

        pass
    return redirect('index')


def ing(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LCR.objects.filter(owner=member).order_by('-id')
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(LCR_hash=hash)

        except:
            hash = None
            process = None

        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len, 'max_index': max_index-2}

        return render(request, 'app/ing.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')



def ing2(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_CI.objects.filter(owner=member).order_by('-id')
        n = len(contract)

        try:
            for i in range(0, n):
                hash = contract.values('sha256')[i]['sha256']
                process = Process.objects.filter(CI_hash=hash)


        except:
            hash = None
            process = None

        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len, 'max_index': max_index - 2}

        return render(request, 'app/ing2.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def ing2_1(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_SR.objects.filter(owner=member).order_by('-id')
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(SR_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len, 'max_index': max_index - 2}

        return render(request, 'app/ing2_1.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def ing3(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LC.objects.filter(owner=member).order_by('-id')

        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(LC_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len, 'max_index': max_index - 2}

        return render(request, 'app/ing3.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def ing4_1(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_BL.objects.filter(owner=member).order_by('-id')

        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(BL_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len, 'max_index': max_index - 2}

        return render(request, 'app/ing4_1.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def ing4_2(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_DO.objects.filter(owner=member).order_by('-id')

        n = len(contract)
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(DO_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len, 'max_index': max_index - 2}

        return render(request, 'app/ing4_2.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def cireceived(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_CI.objects.filter(share1=member).order_by('-id')
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(CI_hash=hash)
        except:
            hash = None
            process = None

        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/cireceived.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')


def srreceived(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_SR.objects.filter(share4=member).order_by('-id')
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(SR_hash=hash)
        except:
            hash = None
            process = None

        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/srreceived.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')


def blreceived1(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_BL.objects.filter(share1=member).order_by('-id')
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(BL_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/blreceived1.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def lcreceived1(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LC.objects.filter(share1=member).order_by('-id')
        n = len(contract)
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(LC_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/lcreceived1.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def blreceived2(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_BL.objects.filter(share2=member).order_by('-id')

        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(BL_hash=hash)
        except:
            hash = None
            process = None

        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/blreceived2.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def lcreceived2(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LC.objects.filter(share2=member).order_by('-id')
        n = len(contract)
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(LC_hash=hash)
        except Exception as e:
            print(e)
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/lcreceived2.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def doreceived(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_DO.objects.filter(share1=member).order_by('-id')
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(DO_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/doreceived.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')

def lcrreceived(request):
    try:
        user_id = request.session['user_id']
        member = Member.objects.get(user_id=user_id)
        contract = Contract_LCR.objects.filter(share3=member).order_by('-id')
        n = len(contract)
        try:
            hash = contract.values('sha256')[0]['sha256']
            process = Process.objects.filter(LCR_hash=hash)
        except:
            hash = None
            process = None
        total_len = len(contract)
        page = request.GET.get('page')
        paginator = Paginator(contract, 6)

        try:
            lines = paginator.page(page)
        except PageNotAnInteger:
            lines = paginator.page(1)
        except EmptyPage:
            lines = paginator.page(paginator.num_pages)
        index = lines.number - 1
        max_index = len(paginator.page_range)
        start_index = index - 2 if index >= 2 else 0
        if index < 2:
            end_index = 3 - start_index
        else:
            end_index = index + 3 if index <= max_index - 3 else max_index
        page_range = list(paginator.page_range[start_index:end_index])

        notice = {'result_list': lines, 'page_range': page_range, 'total_len': total_len,
                  'max_index': max_index - 2}

        return render(request, 'app/lcrreceived.html', notice)
    except Exception as e:
        print(e)
        return redirect('index')


def logout(request):
    try:
        del request.session['user_role']
        del request.session['user_id']
        return redirect('login')
    except:
        return render(request, 'app/login.html',{})



def index(request):

    try:


        user_id = request.session['user_id']
        user_role = request.session['user_role']


        templates = ''
        if user_role == '1':
            templates = 'app/index.html'
        elif user_role == '2':
            templates = 'app/index2.html'
        elif user_role == '3':
            templates = 'app/index3.html'
        elif user_role == '4':
            templates = 'app/index4.html'
        else:
            templates = 'app/login.html'

        return render(request, templates, {'user_id':user_id})
    except Exception as e:
        print(e)
        return redirect('login')

def charts(request):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'}
    res1 = requests.get('https://quotation-api-cdn.dunamu.com/v1/forex/recent?codes=FRX.KRWUSD', headers=headers)
    json_data = res1.json()
    basePrice1 = json_data[0]['basePrice']
    sellprice1 = json_data[0]['cashSellingPrice']
    buyprice1 = json_data[0]['cashBuyingPrice']
    date1=json_data[0]['date']
    time1 = json_data[0]['time']

    res2 = requests.get('https://quotation-api-cdn.dunamu.com/v1/forex/recent?codes=FRX.KRWJPY', headers=headers)
    json_data = res2.json()
    basePrice2 = json_data[0]['basePrice']
    sellprice2 = json_data[0]['cashSellingPrice']
    buyprice2 = json_data[0]['cashBuyingPrice']

    res3 = requests.get('https://quotation-api-cdn.dunamu.com/v1/forex/recent?codes=FRX.KRWCNY', headers=headers)
    json_data = res3.json()
    basePrice3 = json_data[0]['basePrice']
    sellprice3 = json_data[0]['cashSellingPrice']
    buyprice3 = json_data[0]['cashBuyingPrice']

    res4 = requests.get('https://quotation-api-cdn.dunamu.com/v1/forex/recent?codes=FRX.KRWEUR', headers=headers)
    json_data = res4.json()
    basePrice4 = json_data[0]['basePrice']
    sellprice4 = json_data[0]['cashSellingPrice']
    buyprice4 = json_data[0]['cashBuyingPrice']



    try:

        user_role = request.session['user_role']


        templates = ''
        if user_role == '1':
            templates = 'app/charts.html'
        elif user_role == '2':
            templates = 'app/charts2.html'
        elif user_role == '3':
            templates = 'app/charts3.html'
        elif user_role == '4':
            templates = 'app/charts4.html'
        else:
            templates = 'app/login.html'

        return render(request, templates, {'basePrice1': basePrice1,'sellprice1':sellprice1,'buyprice1':buyprice1,'date1':date1,'time1':time1,
                                           'basePrice2':basePrice2,'sellprice2':sellprice2,'buyprice2':buyprice2,
                                           'basePrice3':basePrice3,'sellprice3':sellprice3,'buyprice3':buyprice3,
                                           'basePrice4': basePrice4, 'sellprice4': sellprice4, 'buyprice4': buyprice4})


    except Exception as e:
        print(e)
        return redirect('login')




def forms(request):
    return render(request, 'app/forms.html', {})

def forms2(request):
    return render(request, 'app/forms2.html', {})

def forms2_1(request):
    return render(request, 'app/forms2_1.html', {})

def forms3(request):
    return render(request, 'app/forms3.html', {})

def forms4_1(request):
    return render(request, 'app/forms4_1.html', {})

def forms4_2(request):
    return render(request, 'app/forms4_2.html', {})

def login(request):
    if request.method == 'GET':
        return render(request, 'app/login.html', {})
    else:
        email = request.POST['email']
        password = request.POST['password']
        user_role = request.POST['user_role']

        result_dict = {}
        try:
            Member.objects.get(user_role=user_role, user_id=email, user_pw=password)
            result_dict['result'] = 'success'
            request.session['user_id'] = email
            request.session['user_role'] = user_role
        except Member.DoesNotExist:
            result_dict['result'] = 'fail'
        return JsonResponse(result_dict)


def registerpage(request):
    return render(request, 'app/register.html',{})

def register(request):

    result_dict = {}
    try:
        user_role = request.POST.get('user_role', False)
    except:
        user_role = ''

    user_name = request.POST.get('user_name', False)
    user_id = request.POST.get('user_id', False)
    user_pw = request.POST.get('user_pw', False)
    user_confirm_pw = request.POST.get('user_confirm_pw', False)

    if user_role == '' or user_name == '' or user_id == '' or user_pw == '' or user_confirm_pw == '':
        result_dict['result'] = '공백은 사용할 수 없습니다.'
        return JsonResponse(result_dict)

    elif user_pw != user_confirm_pw:
        result_dict['result'] = '비밀번호 매치 실패'
        return JsonResponse(result_dict)

    else:
        try:
            Member.objects.get(user_id=user_id)
            result_dict['result'] = '이미 가입된 아이디가 있습니다.'
        except Member.DoesNotExist:
            member = Member(user_role=user_role, user_id=user_id, user_pw=user_pw, user_name=user_name)
            member.c_date = timezone.now()
            member.save()
            result_dict['result'] = '가입완료'

        return JsonResponse(result_dict)


def forgot(request):
    return render(request, 'app/forgot.html', {})

